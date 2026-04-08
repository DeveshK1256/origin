from __future__ import annotations

import base64
import re
from datetime import datetime, timezone
from io import BytesIO
from typing import Any

from docx import Document

from services.job_matching import analyze_job_description
from services.nlp_utils import normalize_whitespace


def _safe_str(value: Any, default: str = "") -> str:
    text = str(value or "").strip()
    return text or default


def _coerce_int(value: Any, default: int = 0) -> int:
    try:
        return int(float(str(value).strip()))
    except (TypeError, ValueError):
        return default


def _split_multiline_items(value: Any, limit: int = 8) -> list[str]:
    if isinstance(value, list):
        raw_items = [str(item) for item in value]
    else:
        raw_items = re.split(r"[\r\n]+", str(value or ""))

    cleaned: list[str] = []
    seen: set[str] = set()
    for item in raw_items:
        line = re.sub(r"^\s*[-*]\s*", "", item).strip()
        line = re.sub(r"\s+", " ", line)
        if not line:
            continue
        key = line.lower()
        if key in seen:
            continue
        seen.add(key)
        cleaned.append(line)
        if len(cleaned) >= limit:
            break
    return cleaned


def _split_skill_items(value: Any, limit: int = 20) -> list[str]:
    if isinstance(value, list):
        raw_items = [str(item) for item in value]
    else:
        raw_items = re.split(r"[,\n;/|]+", str(value or ""))

    cleaned: list[str] = []
    seen: set[str] = set()
    for item in raw_items:
        skill = re.sub(r"\s+", " ", item).strip()
        if not skill:
            continue
        key = skill.lower()
        if key in seen:
            continue
        seen.add(key)
        cleaned.append(skill)
        if len(cleaned) >= limit:
            break
    return cleaned


def _profile_from_scratch(scratch_profile: dict[str, Any]) -> dict[str, Any]:
    return {
        "name": _safe_str(scratch_profile.get("name"), "Candidate Name"),
        "email": _safe_str(scratch_profile.get("email")),
        "phone": _safe_str(scratch_profile.get("phone")),
        "location": _safe_str(scratch_profile.get("location")),
        "linkedin": _safe_str(scratch_profile.get("linkedin")),
        "github": _safe_str(scratch_profile.get("github")),
        "summary": _safe_str(scratch_profile.get("summary")),
        "skills": _split_skill_items(scratch_profile.get("skills")),
        "experience": _split_multiline_items(scratch_profile.get("experience_highlights")),
        "education": _split_multiline_items(scratch_profile.get("education")),
        "projects": _split_multiline_items(scratch_profile.get("projects")),
        "certifications": _split_multiline_items(scratch_profile.get("certifications")),
        "achievements": _split_multiline_items(scratch_profile.get("achievements")),
        "experience_years": _coerce_int(scratch_profile.get("experience_years"), 0),
    }


def _profile_from_parsed(resume_parsed: dict[str, Any]) -> dict[str, Any]:
    contact = resume_parsed.get("contact") or {}
    emails = contact.get("emails") or []
    phones = contact.get("phones") or []

    experience_lines: list[str] = []
    experience_lines.extend(_split_multiline_items(resume_parsed.get("recent_roles"), limit=6))
    experience_lines.extend(
        _split_multiline_items((resume_parsed.get("experience_breakdown") or {}).get("evidence"), limit=6)
    )
    experience_lines = _split_multiline_items(experience_lines, limit=8)

    summary = _safe_str(resume_parsed.get("profile_summary"))
    if not summary:
        summary = _safe_str(resume_parsed.get("preview"))
        summary = summary[:300]

    return {
        "name": _safe_str(resume_parsed.get("name"), "Candidate Name"),
        "email": _safe_str(emails[0] if emails else ""),
        "phone": _safe_str(phones[0] if phones else ""),
        "location": _safe_str(resume_parsed.get("location")),
        "linkedin": _safe_str(contact.get("linkedin")),
        "github": _safe_str(contact.get("github")),
        "summary": summary,
        "skills": _split_skill_items(resume_parsed.get("skills")),
        "experience": experience_lines,
        "education": _split_multiline_items(resume_parsed.get("education")),
        "projects": _split_multiline_items(resume_parsed.get("projects")),
        "certifications": _split_multiline_items(resume_parsed.get("certifications")),
        "achievements": _split_multiline_items(resume_parsed.get("achievements")),
        "experience_years": _coerce_int(resume_parsed.get("experience_years"), 0),
    }


def _tailor_profile_to_job(
    profile: dict[str, Any], job_data: dict[str, Any]
) -> tuple[dict[str, Any], dict[str, Any]]:
    required_skills = [skill for skill in (job_data.get("required_skills") or []) if str(skill).strip()]
    preferred_skills = [skill for skill in (job_data.get("preferred_skills") or []) if str(skill).strip()]
    role_title = _safe_str(job_data.get("role_title"), "Target Role")
    role_family = _safe_str(job_data.get("role_family"), "General")
    required_experience = _coerce_int(job_data.get("required_experience_years"), 0)

    available_skills = profile.get("skills") or []
    available_lookup = {skill.lower(): skill for skill in available_skills}

    matched_required = [available_lookup[item.lower()] for item in required_skills if item.lower() in available_lookup]
    matched_preferred = [available_lookup[item.lower()] for item in preferred_skills if item.lower() in available_lookup]
    remaining_skills = [
        item for item in available_skills if item not in matched_required and item not in matched_preferred
    ]

    ordered_skills = (matched_required + matched_preferred + remaining_skills)[:20]
    missing_required = [item for item in required_skills if item.lower() not in available_lookup][:8]

    summary = normalize_whitespace(profile.get("summary") or "")
    summary = summary[:420]
    if summary:
        tailored_summary = (
            f"{summary} Targeting {role_title} opportunities and aligning impact with {role_family} expectations."
        )
    else:
        if matched_required:
            skill_phrase = ", ".join(matched_required[:4])
            tailored_summary = (
                f"Targeting {role_title} opportunities with strengths in {skill_phrase}, "
                "focused on measurable business impact."
            )
        else:
            tailored_summary = (
                f"Targeting {role_title} opportunities with a focus on delivery, collaboration, and measurable outcomes."
            )

    years = _coerce_int(profile.get("experience_years"), 0)
    if required_experience > 0 and years > 0:
        if years >= required_experience:
            tailored_summary += f" Experience level aligns with {required_experience}+ years expected for the role."
        else:
            tailored_summary += (
                f" Current profile highlights transferable outcomes while building toward the {required_experience}+ years expectation."
            )

    tailored_profile = {
        **profile,
        "summary": tailored_summary.strip(),
        "skills": ordered_skills,
    }
    insights = {
        "role_title": role_title,
        "role_family": role_family,
        "required_experience_years": required_experience,
        "matched_required_skills": matched_required[:10],
        "missing_required_skills": missing_required,
    }
    return tailored_profile, insights


def _latex_escape(text: str) -> str:
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    return "".join(replacements.get(char, char) for char in text)


def _latex_list(items: list[str], fallback: str) -> str:
    lines = items or [fallback]
    escaped = [f"\\item {_latex_escape(line)}" for line in lines]
    return "\n".join(escaped)


def _render_latex_resume(profile: dict[str, Any], insights: dict[str, Any]) -> str:
    name = _latex_escape(_safe_str(profile.get("name"), "Candidate Name"))
    email = _safe_str(profile.get("email"))
    phone = _safe_str(profile.get("phone"))
    location = _safe_str(profile.get("location"))
    linkedin = _safe_str(profile.get("linkedin"))
    github = _safe_str(profile.get("github"))
    summary = _latex_escape(_safe_str(profile.get("summary"), ""))

    contact_parts = [part for part in [email, phone, location] if part]
    contact_line = _latex_escape(" | ".join(contact_parts)) if contact_parts else "Contact details available on request"

    links = []
    if linkedin:
        links.append(rf"\href{{{linkedin}}}{{LinkedIn}}")
    if github:
        links.append(rf"\href{{{github}}}{{GitHub}}")
    links_line = " | ".join(links) if links else "Professional links available on request"

    skills = profile.get("skills") or []
    skills_line = _latex_escape(", ".join(skills)) if skills else "Add role-relevant skills"

    experience_items = _latex_list(profile.get("experience") or [], "Add measurable experience bullet points.")
    education_items = _latex_list(profile.get("education") or [], "Add education details.")
    project_items = _latex_list(profile.get("projects") or [], "Add project highlights relevant to target role.")
    cert_items = _latex_list(profile.get("certifications") or [], "Add certifications if applicable.")
    achievement_items = _latex_list(profile.get("achievements") or [], "Add quantified achievements.")

    role_title = _latex_escape(_safe_str(insights.get("role_title"), "Target Role"))
    role_family = _latex_escape(_safe_str(insights.get("role_family"), "General"))
    matched = insights.get("matched_required_skills") or []
    missing = insights.get("missing_required_skills") or []
    matched_line = _latex_escape(", ".join(matched) if matched else "No direct required-skill matches detected yet")
    missing_line = _latex_escape(", ".join(missing) if missing else "No major required-skill gaps detected")

    return (
        "\\documentclass[11pt]{article}\n"
        "\\usepackage[margin=0.85in]{geometry}\n"
        "\\usepackage[hidelinks]{hyperref}\n"
        "\\usepackage{enumitem}\n"
        "\\setlist[itemize]{leftmargin=*,noitemsep,topsep=2pt}\n"
        "\\setlength{\\parindent}{0pt}\n"
        "\\begin{document}\n"
        f"{{\\LARGE \\textbf{{{name}}}}}\\\\\n"
        f"{contact_line}\\\\\n"
        f"{links_line}\n\n"
        "\\section*{Professional Summary}\n"
        f"{summary}\n\n"
        "\\section*{Target Role Alignment}\n"
        f"Role: \\textbf{{{role_title}}} ({role_family})\\\\\n"
        f"Matched required skills: {matched_line}\\\\\n"
        f"Development focus: {missing_line}\n\n"
        "\\section*{Core Skills}\n"
        f"{skills_line}\n\n"
        "\\section*{Experience Highlights}\n"
        "\\begin{itemize}\n"
        f"{experience_items}\n"
        "\\end{itemize}\n\n"
        "\\section*{Projects}\n"
        "\\begin{itemize}\n"
        f"{project_items}\n"
        "\\end{itemize}\n\n"
        "\\section*{Education}\n"
        "\\begin{itemize}\n"
        f"{education_items}\n"
        "\\end{itemize}\n\n"
        "\\section*{Certifications}\n"
        "\\begin{itemize}\n"
        f"{cert_items}\n"
        "\\end{itemize}\n\n"
        "\\section*{Achievements}\n"
        "\\begin{itemize}\n"
        f"{achievement_items}\n"
        "\\end{itemize}\n\n"
        "\\end{document}\n"
    )


def _render_text_resume(profile: dict[str, Any], insights: dict[str, Any]) -> str:
    def block(title: str, items: list[str], fallback: str) -> list[str]:
        lines = [f"=== {title} ==="]
        if items:
            lines.extend(f"- {item}" for item in items)
        else:
            lines.append(f"- {fallback}")
        lines.append("")
        return lines

    contact = " | ".join(
        item for item in [_safe_str(profile.get("email")), _safe_str(profile.get("phone")), _safe_str(profile.get("location"))] if item
    )
    links = " | ".join(item for item in [_safe_str(profile.get("linkedin")), _safe_str(profile.get("github"))] if item)

    lines = [
        "AI Recruitment Intelligence - Resume AI Output",
        f"Generated At: {datetime.now(timezone.utc).isoformat()}",
        "",
        f"Name: {_safe_str(profile.get('name'), 'Candidate Name')}",
        f"Contact: {contact or 'N/A'}",
        f"Links: {links or 'N/A'}",
        "",
        "=== Professional Summary ===",
        _safe_str(profile.get("summary"), "N/A"),
        "",
        "=== Target Role Alignment ===",
        f"- Role: {_safe_str(insights.get('role_title'), 'Target Role')} ({_safe_str(insights.get('role_family'), 'General')})",
        "- Matched required skills: "
        + (", ".join(insights.get("matched_required_skills") or []) or "None detected"),
        "- Development focus: "
        + (", ".join(insights.get("missing_required_skills") or []) or "No major required-skill gaps"),
        "",
    ]

    lines.extend(block("Core Skills", [", ".join(profile.get("skills") or [])] if profile.get("skills") else [], "Add role-relevant skills."))
    lines.extend(block("Experience Highlights", profile.get("experience") or [], "Add measurable experience bullet points."))
    lines.extend(block("Projects", profile.get("projects") or [], "Add project highlights relevant to target role."))
    lines.extend(block("Education", profile.get("education") or [], "Add education details."))
    lines.extend(block("Certifications", profile.get("certifications") or [], "Add certifications if applicable."))
    lines.extend(block("Achievements", profile.get("achievements") or [], "Add quantified achievements."))
    return "\n".join(lines).strip()


def _render_docx_resume(profile: dict[str, Any], insights: dict[str, Any]) -> bytes:
    doc = Document()
    doc.add_heading(_safe_str(profile.get("name"), "Candidate Name"), level=0)

    contact = " | ".join(
        item for item in [_safe_str(profile.get("email")), _safe_str(profile.get("phone")), _safe_str(profile.get("location"))] if item
    )
    if contact:
        doc.add_paragraph(contact)

    links = " | ".join(item for item in [_safe_str(profile.get("linkedin")), _safe_str(profile.get("github"))] if item)
    if links:
        doc.add_paragraph(links)

    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(_safe_str(profile.get("summary"), "N/A"))

    doc.add_heading("Target Role Alignment", level=1)
    doc.add_paragraph(
        f"Role: {_safe_str(insights.get('role_title'), 'Target Role')} ({_safe_str(insights.get('role_family'), 'General')})"
    )
    doc.add_paragraph(
        "Matched required skills: " + (", ".join(insights.get("matched_required_skills") or []) or "None detected")
    )
    doc.add_paragraph(
        "Development focus: " + (", ".join(insights.get("missing_required_skills") or []) or "No major required-skill gaps")
    )

    doc.add_heading("Core Skills", level=1)
    doc.add_paragraph(", ".join(profile.get("skills") or []) or "Add role-relevant skills")

    sections = [
        ("Experience Highlights", profile.get("experience") or [], "Add measurable experience bullet points."),
        ("Projects", profile.get("projects") or [], "Add project highlights relevant to target role."),
        ("Education", profile.get("education") or [], "Add education details."),
        ("Certifications", profile.get("certifications") or [], "Add certifications if applicable."),
        ("Achievements", profile.get("achievements") or [], "Add quantified achievements."),
    ]
    for title, items, fallback in sections:
        doc.add_heading(title, level=1)
        for line in (items or [fallback]):
            doc.add_paragraph(line, style="List Bullet")

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def generate_resume_ai_assets(
    source_mode: str,
    job_description: str,
    scratch_profile: dict[str, Any] | None = None,
    resume_parsed: dict[str, Any] | None = None,
) -> dict[str, Any]:
    normalized_mode = _safe_str(source_mode, "scratch").lower()
    cleaned_job_description = normalize_whitespace(job_description)
    if not cleaned_job_description:
        raise ValueError("Job description is required for resume generation.")

    job_data = analyze_job_description(cleaned_job_description)

    if normalized_mode == "existing":
        if not isinstance(resume_parsed, dict) or not resume_parsed:
            raise ValueError("For existing mode, provide parsed resume data.")
        profile = _profile_from_parsed(resume_parsed)
    else:
        profile = _profile_from_scratch(scratch_profile or {})

    tailored_profile, insights = _tailor_profile_to_job(profile, job_data)
    latex_code = _render_latex_resume(tailored_profile, insights)
    resume_text = _render_text_resume(tailored_profile, insights)
    resume_docx_bytes = _render_docx_resume(tailored_profile, insights)

    stamp = datetime.now(timezone.utc).strftime("%Y%m%d-%H%M%S")
    return {
        "source_mode": normalized_mode,
        "latex_code": latex_code,
        "resume_text": resume_text,
        "resume_docx_base64": base64.b64encode(resume_docx_bytes).decode("ascii"),
        "latex_filename": f"tailored-resume-{stamp}.tex",
        "resume_filename": f"tailored-resume-{stamp}.docx",
        "overleaf_url": "https://www.overleaf.com/docs",
        "overleaf_hint": "Open Overleaf, create a blank project, and paste the generated LaTeX code.",
        "job_target": {
            "role_title": insights.get("role_title"),
            "role_family": insights.get("role_family"),
            "required_experience_years": insights.get("required_experience_years"),
            "matched_required_skills": insights.get("matched_required_skills"),
            "missing_required_skills": insights.get("missing_required_skills"),
        },
    }
