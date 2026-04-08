from __future__ import annotations

import os
import zipfile
from xml.etree import ElementTree
from io import BytesIO

from docx import Document
from PyPDF2 import PdfReader

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}


def allowed_resume_extension(filename: str) -> bool:
    _, ext = os.path.splitext(filename.lower())
    return ext in ALLOWED_EXTENSIONS


def _clean_extracted_text(text: str) -> str:
    lines = [line.strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line)


def _extract_from_pdf_with_pdfplumber(binary_data: bytes) -> str:
    try:
        import pdfplumber  # type: ignore
    except Exception:
        return ""

    parts: list[str] = []
    with pdfplumber.open(BytesIO(binary_data)) as pdf:
        for page in pdf.pages:
            content = page.extract_text() or ""
            if content.strip():
                parts.append(content)
    return _clean_extracted_text("\n".join(parts))


def _extract_from_pdf(binary_data: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(binary_data), strict=False)
    except Exception as exc:
        raise ValueError(f"Could not open PDF file: {exc}") from exc

    if reader.is_encrypted:
        try:
            decrypt_status = reader.decrypt("")
            if decrypt_status == 0:
                raise ValueError("Encrypted PDF is not supported.")
        except Exception as exc:
            raise ValueError(f"Encrypted or protected PDF cannot be parsed: {exc}") from exc

    parts: list[str] = []
    for page in reader.pages:
        try:
            content = page.extract_text() or ""
        except Exception:
            content = ""
        if content.strip():
            parts.append(content)

    text = _clean_extracted_text("\n".join(parts))
    if text:
        return text

    fallback = _extract_from_pdf_with_pdfplumber(binary_data)
    if fallback:
        return fallback

    raise ValueError(
        "No readable text found in PDF. The file may be scanned/image-only or protected."
    )


def _extract_docx_with_python_docx(binary_data: bytes) -> str:
    document = Document(BytesIO(binary_data))
    lines: list[str] = []

    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    lines.append(text)

    return _clean_extracted_text("\n".join(lines))


def _extract_docx_with_xml(binary_data: bytes) -> str:
    try:
        with zipfile.ZipFile(BytesIO(binary_data)) as archive:
            xml_content = archive.read("word/document.xml")
    except Exception:
        return ""

    try:
        root = ElementTree.fromstring(xml_content)
    except ElementTree.ParseError:
        return ""

    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    texts = [node.text.strip() for node in root.findall(".//w:t", namespace) if node.text and node.text.strip()]
    return _clean_extracted_text("\n".join(texts))


def _extract_from_docx(binary_data: bytes) -> str:
    try:
        parsed = _extract_docx_with_python_docx(binary_data)
        if parsed:
            return parsed
    except Exception:
        pass

    fallback = _extract_docx_with_xml(binary_data)
    if fallback:
        return fallback

    raise ValueError("Could not parse DOCX file content.")


def _extract_from_txt(binary_data: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            decoded = binary_data.decode(encoding)
            if decoded.strip():
                return decoded
        except Exception:
            continue
    return binary_data.decode("utf-8", errors="ignore")


def extract_text_from_uploaded_file(file_storage) -> str:
    filename = file_storage.filename or ""
    _, ext = os.path.splitext(filename.lower())
    binary_data = file_storage.read()

    if not binary_data:
        raise ValueError("Uploaded file contains no data.")

    if ext == ".pdf":
        text = _extract_from_pdf(binary_data)
    elif ext == ".docx":
        text = _extract_from_docx(binary_data)
    elif ext == ".txt":
        text = _extract_from_txt(binary_data)
    else:
        raise ValueError("Unsupported file type. Please upload PDF, DOCX, or TXT.")

    normalized = text.strip()
    if not normalized:
        raise ValueError(
            "File parsed but no readable text was extracted. Use text-based PDF/DOCX files."
        )
    return normalized
